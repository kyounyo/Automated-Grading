import os
import re
from pathlib import Path
from typing import List, Dict, Any


def extract_text_from_file(file_path: str) -> str:
    """
    Extracts plain text from PDF, DOCX, XLSX, CSV, or TXT files.
    """
    path = Path(file_path)
    if not path.exists():
        raise FileNotFoundError(f"File not found: {file_path}")

    ext = path.suffix.lower()
    
    if ext == ".pdf":
        return _extract_pdf(file_path)
    elif ext in [".docx", ".doc"]:
        return _extract_docx(file_path)
    elif ext in [".xlsx", ".xls", ".csv"]:
        return _extract_excel(file_path)
    elif ext in [".txt", ".md"]:
        with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
            return f.read()
    else:
        with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
            return f.read()


def calculate_question_max_mark(prompt_text: str) -> float:
    """
    Precision extractor for question max marks from PDF & DOCX prompt text ONLY:
    1. Looks for "(Total: X marks)", "[Total X Marks]", "Total: X marks".
    2. Sums parenthesized sub-part marks e.g. "(5 marks)", "(2 marks)", "[5 pts]".
    3. Handles multiplier patterns like "(5 marks each)".
    4. Filters out sub-part labels like (1), (a).
    """
    if not prompt_text:
        return 10.0

    prompt_only = re.split(r'Marking\s+Rubric|Answer\s+Scheme|Model\s+Answer|Marking\s+Scheme|Rubric', prompt_text, flags=re.IGNORECASE)[0].strip()

    # 1. Total mark explicitly stated e.g. "(Total: 25 marks)", "[Total 20 Marks]", "Total: 10 marks"
    total_m = re.search(r'(?:total\s*:?\s*)(\d+)\s*(?:marks?|pts?|points?)?', prompt_only, re.IGNORECASE)
    if total_m:
        val = float(total_m.group(1))
        if val > 0: return val

    # 2. Sum parenthesized / bracketed sub-part marks e.g. "(5 marks)", "[2 pts]"
    part_marks = re.findall(r'[\(\[]\s*(\d+)\s*(?:marks?|pts?|points?)\s*[\)\]]', prompt_only, re.IGNORECASE)
    if part_marks:
        explicit_marks = [int(m) for m in part_marks if re.search(r'[\(\[]\s*' + m + r'\s*(?:marks?|pts?|points?)\s*[\)\]]', prompt_only, re.IGNORECASE)]
        if explicit_marks:
            return float(sum(explicit_marks))

    # 3. Multiplier patterns e.g. "(5 marks each)" or "5 marks each"
    each_m = re.search(r'(\d+)\s*marks?\s*each', prompt_only, re.IGNORECASE)
    if each_m:
        per_part = int(each_m.group(1))
        parts_count = len(re.findall(r'\([a-z0-9]+\)', prompt_only, re.IGNORECASE))
        if parts_count > 0:
            return float(per_part * parts_count)

    # 4. Trailing "X marks" pattern
    end_marks = re.findall(r'\b(\d+)\s*(?:marks?|pts?|points?)\b', prompt_only, re.IGNORECASE)
    if end_marks:
        total = sum(int(m) for m in end_marks)
        if total > 0: return float(total)

    return 10.0


def parse_excel_rubric(file_path: str) -> List[Dict[str, Any]]:
    """
    Template-tolerant parser for Excel/CSV rubric schemes.
    Uses header detection, alias matching, and content validation.
    Falls back to positional heuristic parser for backward compatibility.
    """
    try:
        from .flexible_excel_parser import parse_flexible_rubric
        parsed = parse_flexible_rubric(file_path)
        if parsed and len(parsed) > 0:
            return parsed
    except Exception as e:
        print(f"[DocumentParser Warning] Flexible rubric parser exception: {e}, attempting fallback.")

    # Deterministic Legacy Fallback
    try:
        import pandas as pd
        ext = Path(file_path).suffix.lower()
        if ext == ".csv":
            df = pd.read_csv(file_path)
        else:
            df = pd.read_excel(file_path)

        cols = [str(c).strip().lower() for c in df.columns]
        
        q_num_col = next((df.columns[i] for i, c in enumerate(cols) if any(k in c for k in ["question_n", "question_no", "q_num", "q_no", "num"])), None)
        q_text_col = next((df.columns[i] for i, c in enumerate(cols) if c in ["question", "prompt", "question_text", "criteria"] or ("question" in c and c != q_num_col)), None)
        ans_col = next((df.columns[i] for i, c in enumerate(cols) if any(k in c for k in ["answer", "rubric", "model_answer", "marking", "solution"])), None)
        mark_col = next((df.columns[i] for i, c in enumerate(cols) if any(k in c for k in ["max_mark", "mark", "score", "max_score", "points"])), None)

        questions = []
        for idx, row in df.iterrows():
            q_num_val = str(row[q_num_col]).strip() if q_num_col and pd.notna(row[q_num_col]) else str(idx + 1)
            if q_num_val.endswith('.0'):
                q_num_val = q_num_val[:-2]

            q_text_val = str(row[q_text_col]).strip() if q_text_col and pd.notna(row[q_text_col]) else f"Question {q_num_val}"
            ans_val = str(row[ans_col]).strip() if ans_col and pd.notna(row[ans_col]) else q_text_val
            
            try:
                max_mark_val = float(row[mark_col]) if mark_col and pd.notna(row[mark_col]) else 10.0
            except Exception:
                max_mark_val = 10.0

            q_label = q_num_val if q_num_val.lower().startswith("q") else f"Q{q_num_val}"

            questions.append({
                "id": idx + 1,
                "question_number": q_label,
                "text": q_text_val,
                "maxMark": max_mark_val if max_mark_val > 0 else 10.0,
                "modelAnswer": ans_val
            })

        return questions
    except Exception as e:
        print(f"[DocumentParser Error] Excel rubric parsing fallback failed on {file_path}: {e}")
        return []


def parse_separate_question_and_rubric_docs(q_doc: str, r_doc: str) -> List[Dict[str, Any]]:
    """
    Pairs a Question Paper document with a separate Marking Rubric document for DOCX & PDF files.
    """
    q_matches = list(re.finditer(r"(?:^|\n|\s{2,})(?:Question|Q)?\s*(\d+)[\.\s]", q_doc, re.IGNORECASE))
    if not q_matches:
        q_matches = list(re.finditer(r"(?:Question|Q)\s*(\d+)", q_doc, re.IGNORECASE))
        
    parsed = []
    seen = set()

    for i in range(len(q_matches)):
        q_num = q_matches[i].group(1)
        if q_num in seen:
            continue
        seen.add(q_num)

        start = q_matches[i].start()
        next_start = len(q_doc)
        for j in range(i + 1, len(q_matches)):
            if q_matches[j].group(1) != q_num:
                next_start = q_matches[j].start()
                break

        prompt_text = q_doc[start:next_start].strip()

        # Find exact "Marking Rubric for Question X" section in r_doc
        m_rubric = re.search(r"Marking\s+(?:Rubric\s+)?for\s+(?:Question|Q)?\s*" + q_num, r_doc, re.IGNORECASE)
        if m_rubric:
            r_start = m_rubric.start()
            r_all = list(re.finditer(r"Marking\s+(?:Rubric\s+)?for\s+(?:Question|Q)?\s*\d+", r_doc, re.IGNORECASE))
            r_end = len(r_doc)
            for rm in r_all:
                if rm.start() > r_start:
                    r_end = rm.start()
                    break
            rubric_text = r_doc[r_start:r_end].strip()
        else:
            rubric_text = prompt_text

        total_marks = calculate_question_max_mark(prompt_text)

        parsed.append({
            "id": len(parsed) + 1,
            "question_number": f"Q{q_num}",
            "text": prompt_text,
            "maxMark": total_marks,
            "modelAnswer": rubric_text
        })

    return parsed


def smart_parse_rubric_text(raw_text: str) -> List[Dict[str, Any]]:
    """
    Intelligently splits raw PDF/Docx text into distinct Question blocks (e.g. Question 6, Question 8).
    """
    clean_text = re.sub(r'--- Page \d+ ---', '', raw_text)
    clean_text = re.sub(r'[\u2060\u200b\ufeff]', '', clean_text)
    clean_text = re.sub(r'\n{3,}', '\n\n', clean_text).strip()

    q_matches = list(re.finditer(r'(?<!for\s)(?<!Rubric\s)(?<!Marking\s)\b(?:Question|Q)\s*(\d+)', clean_text, re.IGNORECASE))
    if not q_matches:
        q_matches = list(re.finditer(r'(?:^|\n)\s*(\d+)[\.\)]', clean_text))

    if not q_matches:
        return []

    parsed = []
    seen_nums = set()

    for i in range(len(q_matches)):
        q_num = q_matches[i].group(1)
        if q_num in seen_nums:
            continue
        seen_nums.add(q_num)

        start_idx = q_matches[i].start()
        next_start = len(clean_text)
        for j in range(i + 1, len(q_matches)):
            if q_matches[j].group(1) != q_num:
                next_start = q_matches[j].start()
                break

        block = clean_text[start_idx:next_start].strip()

        rubric_match = re.search(r'Marking\s+(?:Rubric\s+)?for\s+(?:Question|Q)?\s*' + q_num, block, re.IGNORECASE)
        if rubric_match:
            prompt_text = block[:rubric_match.start()].strip()
            rubric_text = block[rubric_match.start():].strip()
        else:
            prompt_text = block
            rubric_text = block

        total_marks = calculate_question_max_mark(prompt_text)

        parsed.append({
            "id": len(parsed) + 1,
            "question_number": f"Q{q_num}",
            "text": prompt_text,
            "maxMark": total_marks,
            "modelAnswer": rubric_text
        })

    return parsed


def parse_excel_rows(file_path: str) -> List[Dict[str, Any]]:
    """
    Template-tolerant parser for Excel/CSV student submission datasets.
    Uses header detection, alias matching, and content validation.
    Falls back to legacy heuristic parser for backward compatibility.
    """
    try:
        from .flexible_excel_parser import parse_flexible_submissions
        parsed = parse_flexible_submissions(file_path)
        if parsed and len(parsed) > 0:
            return parsed
    except Exception as e:
        print(f"[DocumentParser Warning] Flexible submissions parser exception: {e}, attempting fallback.")

    # Deterministic Legacy Fallback
    try:
        import pandas as pd
        ext = Path(file_path).suffix.lower()
        if ext == ".csv":
            df = pd.read_csv(file_path)
        else:
            df = pd.read_excel(file_path)

        cols = [str(c).strip().lower() for c in df.columns]
        
        stu_col = next((df.columns[i] for i, c in enumerate(cols) if any(k in c for k in ["student_id", "student_no", "stu_id", "student", "id"])), df.columns[0])
        email_col = next((df.columns[i] for i, c in enumerate(cols) if any(k in c for k in ["student_gmail", "student_gma", "student_email", "gmail", "email", "gma", "mail"])), None)
        name_col = next((df.columns[i] for i, c in enumerate(cols) if any(k in c for k in ["student_name", "student_nam", "name", "nam"])), None)
        q_col = next((df.columns[i] for i, c in enumerate(cols) if any(k in c for k in ["question_no", "question_n", "question", "q_no", "q_num"])), None)
        resp_col = next((df.columns[i] for i, c in enumerate(cols) if any(k in c for k in ["response", "answer", "student_answer", "submission", "text"])), df.columns[-1])

        # Group by Student ID if question_no column exists
        if q_col:
            students = {}
            for idx, row in df.iterrows():
                s_id = str(row[stu_col]).strip() if pd.notna(row[stu_col]) else f"STU{1000 + idx}"
                if s_id.endswith(".0"): s_id = s_id[:-2]

                s_name = str(row[name_col]).strip() if (name_col and pd.notna(row[name_col]) and str(row[name_col]).strip() and str(row[name_col]).strip() != "nan") else f"Student {s_id}"
                s_email = str(row[email_col]).strip() if (email_col and pd.notna(row[email_col]) and str(row[email_col]).strip() and str(row[email_col]).strip() != "nan") else "N/A"

                q_num = str(row[q_col]).strip() if pd.notna(row[q_col]) else f"Q{idx + 1}"
                if q_num.endswith(".0"): q_num = q_num[:-2]
                q_label = q_num if q_num.lower().startswith("q") else f"Q{q_num}"

                resp_text = str(row[resp_col]).strip() if pd.notna(row[resp_col]) else ""

                if s_id not in students:
                    students[s_id] = {
                        "name": s_name,
                        "email": s_email,
                        "responses": []
                    }
                else:
                    if students[s_id]["name"].startswith("Student ") and not s_name.startswith("Student "):
                        students[s_id]["name"] = s_name
                    if students[s_id]["email"] == "N/A" and s_email != "N/A":
                        students[s_id]["email"] = s_email

                students[s_id]["responses"].append(f"Question {q_label}:\n{resp_text}")

            rows = []
            for s_id, s_data in students.items():
                rows.append({
                    "student_id": s_id,
                    "student_name": s_data["name"],
                    "student_email": s_data["email"],
                    "text": "\n\n".join(s_data["responses"])
                })
            return rows

        # Fallback for single-row per student format
        rows = []
        for idx, row in df.iterrows():
            row_dict = {str(k).strip().lower(): str(v).strip() for k, v in row.items() if pd.notna(v) and str(v).strip()}
            
            student_id = next((str(row_dict[k]) for k in row_dict if "id" in k or "student" in k or "num" in k), f"STU{1000 + idx}")
            student_name = next((str(row_dict[k]) for k in row_dict if "name" in k or "nam" in k), "N/A")
            student_email = next((str(row_dict[k]) for k in row_dict if "email" in k or "gmail" in k or "gma" in k or "mail" in k), "N/A")
            
            text_parts = [f"{k.capitalize()}: {v}" for k, v in row_dict.items() if not any(x in k for x in ["student_id", "id", "student_name", "name", "email", "gmail", "gma", "mail"])]
            full_text = "\n".join(text_parts) if text_parts else str(row.to_dict())

            rows.append({
                "student_id": student_id,
                "student_name": student_name,
                "student_email": student_email,
                "text": full_text
            })
        return rows
    except Exception as e:
        print(f"[DocumentParser Error] Excel row parsing fallback failed: {e}")
        return []


def _extract_pdf(file_path: str) -> str:
    try:
        from pypdf import PdfReader
        reader = PdfReader(file_path)
        extracted = []
        for i, page in enumerate(reader.pages):
            text = page.extract_text()
            if text:
                text = re.sub(r'[\u2060\u200b\ufeff]', '', text)
                extracted.append(text)
        
        full = "\n\n".join(extracted)
        return re.sub(r'\n{3,}', '\n\n', full).strip()
    except Exception as e:
        print(f"[DocumentParser Error] PyPDF failed on {file_path}: {e}")
        return ""


def _extract_docx(file_path: str) -> str:
    try:
        import docx
        doc = docx.Document(file_path)
        extracted = []
        
        # Extract paragraph text
        for p in doc.paragraphs:
            if p.text.strip():
                extracted.append(p.text.strip())
                
        # Extract Word table cell text if present
        for table in doc.tables:
            for row in table.rows:
                row_cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if row_cells:
                    extracted.append(" | ".join(row_cells))

        full_txt = "\n".join(extracted)
        return re.sub(r'\n{3,}', '\n\n', full_txt).strip()
    except Exception as e:
        print(f"[DocumentParser Error] docx parsing failed on {file_path}: {e}")
        return ""


def _extract_excel(file_path: str) -> str:
    try:
        import pandas as pd
        ext = Path(file_path).suffix.lower()
        if ext == ".csv":
            df = pd.read_csv(file_path)
        else:
            df = pd.read_excel(file_path)
        return df.to_string()
    except Exception as e:
        print(f"[DocumentParser Error] excel parsing failed on {file_path}: {e}")
        return ""
